import streamlit as st
import pandas as pd
from utils.call_aws_api import call_api
from utils.load_amazon_reviews import load_data
from docx import Document
from nltk.tokenize import sent_tokenize, word_tokenize 
import nltk
import re

# Download required NLTK data
nltk.download('punkt_tab')

def remove_timestamps(text):
    # Remove patterns like "[Name] 00:00:00" or "[inaudible 00:00:00]"
    text = re.sub(r'\[[^\]]+\]\s*\d{2}:\d{2}:\d{2}|\[[^\]]+\d{2}:\d{2}:\d{2}\]', ' ', text)
    # Remove any remaining timestamps
    text = re.sub(r'\d{2}:\d{2}:\d{2}', ' ', text)
    return text.strip()

def split_into_paragraphs(text, sentence_min_len=8, sentence_max_len=200):
    # Simple sentence splitting on common punctuation
    sentences = sent_tokenize(text)
    
    filtered_sentences = []
    discarded_sentences = []
    for sentence in sentences:
        sentence = sentence.strip().replace('\n', ' ')
        wc = len(word_tokenize(sentence))
        if sentence_min_len <= wc <= sentence_max_len:
            filtered_sentences.append(sentence)
        else:
            reason = "Too short" if wc < sentence_min_len else "Too long"
            discarded_sentences.append({
                "Sentence": sentence,
                "Word Count": wc,
                "Reason": reason
            })
    
    return filtered_sentences, discarded_sentences
# Constants
MAX_REVIEWS = 1000
MAX_CHARS_PER_REVIEW = 2000

def validate_review_length(review: str) -> str:
    """Validate and truncate review if it exceeds maximum length."""
    if len(review) > MAX_CHARS_PER_REVIEW:
        return review[:MAX_CHARS_PER_REVIEW]
    return review

def validate_reviews_count(reviews: list) -> list:
    """Validate and limit the number of reviews."""
    if len(reviews) > MAX_REVIEWS:
        st.warning(f"Currently, we only support up to {MAX_REVIEWS} sentences. Truncating to first {MAX_REVIEWS} sentences.")
        return reviews[:MAX_REVIEWS]
    return reviews

def process_text_document(content: str, sentence_min_len: int, sentence_max_len: int) -> tuple[list, list]:
    """Process text content and return filtered and discarded sentences."""
    content = remove_timestamps(content)
    return split_into_paragraphs(content, sentence_min_len, sentence_max_len)

def process_text_or_doc_file(upload_file, sentence_min_len: int, sentence_max_len: int) -> tuple[pd.DataFrame, str]:
    """Process txt or docx file and return dataframe and column name."""
    input_column = 'Sentences'
    file_type = upload_file.name.lower().split('.')[-1]
    
    if file_type == "txt":
        content = upload_file.read().decode()
    elif file_type == "docx":
        doc = Document(upload_file)
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    sentences, discarded = process_text_document(content, sentence_min_len, sentence_max_len)
    return pd.DataFrame({input_column: sentences}), input_column, discarded

def process_tabular_file(upload_file) -> tuple[pd.DataFrame, str]:
    """Process csv or xlsx file and return dataframe and column name."""
    file_type = upload_file.name.lower().split('.')[-1]
    
    if file_type == "csv":
        df = pd.read_csv(upload_file)
    elif file_type == "xlsx":
        df = pd.read_excel(upload_file)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    input_column = df.columns[0]
    if len(df.columns) > 1:
        input_column = st.selectbox("Select Review Column:", df.columns.tolist())
    
    # Clean timestamps from the selected column
    df[input_column] = df[input_column].apply(remove_timestamps)
    return df, input_column

def display_results(sentences_df: pd.DataFrame, discarded_df: pd.DataFrame | None, output_filename: str):
    """Display and provide download options for results."""
    if discarded_df is not None and not discarded_df.empty:
        st.subheader("Discarded Sentences")
        st.dataframe(discarded_df, use_container_width=True)
    
    st.subheader("Validated Sentences")
    st.dataframe(sentences_df, use_container_width=True)
    
    csv_validated = sentences_df.to_csv(index=False).encode('utf-8')
    st.download_button(
        label="Download CSV of Validated Sentences",
        data=csv_validated,
        file_name=output_filename,
        mime="text/csv"
    )

def process_reviews(input_df: pd.DataFrame, input_column: str, category_name: str, df_placeholder: st.empty, header_placeholder: st.empty) -> pd.DataFrame:
    """Process reviews through the LLM and return results."""
    reviews = input_df[input_column].tolist()
    reviews = validate_reviews_count(reviews)
    
    # Validate review lengths
    valid_reviews = [validate_review_length(review) for review in reviews]
    input_df = input_df.copy()
    input_df[input_column] = valid_reviews
    
    output_df = input_df.copy()
    output_df['LLM Response'] = ''
    
    header_placeholder.subheader("Input & Results")
    progress_text = st.empty()
    progress_bar = st.progress(0)
    status_container = st.empty()
    
    # Add pause button
    stop_button = st.empty()
    is_stopped = stop_button.button("Stop Processing")
    
    total = len(valid_reviews)
    for i, row in input_df.iterrows():
        if is_stopped:
            st.warning("Processing stopped. Partial results are shown below and can be downloaded.")
            break
        
        progress_text.text(f"Processing review {i+1}/{total} (This may take a few minutes)")
        progress_bar.progress((i+1)/total)
        review = row[input_column]
        try:
            response = call_api(category_name, review, status_container=status_container)
        except TimeoutError:
            st.error("Max retries reached. The server takes unexpected long to load. Please try again.")
            break
        except RuntimeError as e:
            st.error(f"Error processing review {i+1}: {str(e)}")
            raise e
        output_df.at[i, 'LLM Response'] = response
        df_placeholder.dataframe(output_df, use_container_width=True)
    
    progress_text.empty()
    status_container.empty()
    progress_bar.empty()
    stop_button.empty()
    return output_df

def handle_file_conversion():
    """Handle the file conversion tab functionality."""
    st.header("Convert Text to CSV of Sentences")
    
    # Initialize session state for storing file content
    if 'file_content' not in st.session_state:
        st.session_state.file_content = None
    if 'text_input' not in st.session_state:
        st.session_state.text_input = None
    
    upload_file = st.file_uploader(
        "Upload *.txt or *.docx File (Optional)", 
        type=["txt", "docx"],
        help="Optional. Upload a text file (.txt) or a Word document (.docx)",
        key="file_uploader"
    )
    
    col1, col2, col3 = st.columns(3)
    with col1:
        output_filename = st.text_input("Output CSV File Name (Optional)", value="converted_file.csv", help="Enter the desired name for your output CSV file")
        if not output_filename.endswith(".csv"):
            output_filename += ".csv"
    with col2:
        sentence_min_len = st.number_input("Min Words per Sentence (Optional)", value=4, min_value=1)
    with col3:
        sentence_max_len = st.number_input("Max Words per Sentence (Optional)", value=200, min_value=1)
    
    if upload_file is not None:
        file_type = upload_file.name.lower().split('.')[-1]
        
        try:
            # Store file content in session state if it's a new file
            if st.session_state.file_content is None or upload_file != st.session_state.last_uploaded_file:
                if file_type == "txt":
                    st.session_state.file_content = upload_file.read().decode()
                elif file_type == "docx":
                    doc = Document(upload_file)
                    st.session_state.file_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                st.session_state.last_uploaded_file = upload_file
            
            if file_type in ["txt", "docx"]:
                sentences, discarded = process_text_document(st.session_state.file_content, sentence_min_len, sentence_max_len)
                sentences_df = pd.DataFrame({"Input": sentences})
                discarded_df = pd.DataFrame(discarded) if discarded else None
                display_results(sentences_df, discarded_df, output_filename)
            
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
    
    else:
        text_input = st.text_area(
            "Enter Text (Optional, not needed if file uploaded)",
            help="Enter text that will be split into sentences. Not needed if you uploaded a file above.",
            key="text_area"
        )
        
        # Store text input in session state if it changes
        if text_input != st.session_state.text_input:
            st.session_state.text_input = text_input
        
        if text_input:
            sentences, discarded = process_text_document(text_input, sentence_min_len, sentence_max_len)
            sentences_df = pd.DataFrame({"Input": sentences})
            discarded_df = pd.DataFrame(discarded) if discarded else None
            display_results(sentences_df, discarded_df, output_filename)

def handle_llm_submission():
    """Handle the LLM submission tab functionality."""
    st.header(f"Submit Sentences to LLM (at most {MAX_REVIEWS} each time)")
    
    category_name = st.text_input(
        "Product/Service Category of the Sentences (Press Enter↵ to Confirm)*",
        help="Required. Enter the category that best describes these sentences (e.g. Electronics, Books, etc.)"
    )
    
    upload_file = st.file_uploader(
        "Upload File (CSV or Excel) *", 
        type=["csv", "xlsx"],
        help="Required. Upload a .csv file or an Excel file with headers."
    )
    
    if upload_file is not None:
        input_df, input_column = process_tabular_file(upload_file)
        display_container = st.container()
        header_placeholder = display_container.empty()
        df_placeholder = display_container.empty()
        with display_container:
            header_placeholder.subheader("Input")
            df_placeholder.dataframe(input_df, use_container_width=True)
            if not category_name:
                st.warning("Please set a Product/Service Category to enable submitting the sentences to LLM.")
            if st.button(f"Submit Sentences to LLM (at most {MAX_REVIEWS} each time)", disabled=not category_name):
                output_df = process_reviews(input_df, input_column, category_name, df_placeholder, header_placeholder)
                csv_response = output_df.to_csv(index=False).encode('utf-8')
                output_filename = st.text_input("LLM Response CSV File Name (Optional)", value=f"{category_name}-response.csv", help="Enter the desired name for your output CSV file")
                st.download_button("Download LLM Response CSV", csv_response, output_filename, "text/csv")

def handle_download_reviews():
    """Handle the download reviews tab functionality."""
    st.header("Download Sample Amazon Reviews from Hugging Face")
    AMAZON_CATEGORIES = [
        "All_Beauty", "Toys_and_Games", "Cell_Phones_and_Accessories", 
        "Industrial_and_Scientific", "Gift_Cards", "Musical_Instruments",
        "Electronics", "Handmade_Products", "Arts_Crafts_and_Sewing",
        "Baby_Products", "Health_and_Household", "Office_Products",
        "Digital_Music", "Grocery_and_Gourmet_Food", "Sports_and_Outdoors", 
        "Home_and_Kitchen", "Subscription_Boxes", "Tools_and_Home_Improvement",
        "Pet_Supplies", "Video_Games", "Kindle_Store",
        "Clothing_Shoes_and_Jewelry", "Patio_Lawn_and_Garden", "Unknown",
        "Books", "Automotive", "CDs_and_Vinyl", "Beauty_and_Personal_Care",
        "Amazon_Fashion", "Magazine_Subscriptions", "Software",
        "Health_and_Personal_Care", "Appliances", "Movies_and_TV"
    ]

    category = st.selectbox(
        "Amazon Review Category *",
        options=AMAZON_CATEGORIES,
        index=AMAZON_CATEGORIES.index("Digital_Music"),
        help="Required. Select the category from the Hugging Face dataset: https://huggingface.co/datasets/McAuley-Lab/Amazon-Reviews-2023#grouped-by-category"
    )
    num_examples = st.number_input(
        "Number of Examples *", 
        min_value=1, 
        max_value=MAX_REVIEWS, 
        value=10,
        help=f"Required. Choose between 1-{MAX_REVIEWS} reviews to download."
    )
    
    if st.button("Load Data"):
        progress_text = st.empty()
        progress_bar = st.progress(0)
        data_container = st.empty()
        
        data = load_data(category, num_examples, progress_text, progress_bar, data_container)
        
        progress_text.empty()
        progress_bar.empty()
        
        if data is not None and not data.empty:
            csv = data.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="Download CSV",
                data=csv,
                file_name=f"{category}-first_{num_examples}_reviews.csv",
                mime="text/csv"
            )

def main():
    """Main function to run the Streamlit app."""
    st.title("Voice of Customer Analysis with Supervised-Finetuned LLM")
    
    tab1, tab2, tab3 = st.tabs(["Convert Text to CSV of Sentences", "Submit to LLM", "Download Amazon Reviews"])
    
    with tab1:
        handle_file_conversion()
    
    with tab2:
        handle_llm_submission()
        
    with tab3:
        handle_download_reviews()

if __name__ == "__main__":
    main()